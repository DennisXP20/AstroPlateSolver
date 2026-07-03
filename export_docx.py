"""
Word-Dokument-Export für den Plate Solver.
Erzeugt einen Bericht mit Bild-Ausschnitt und Daten für jedes erkannte Objekt.

Abhängigkeiten: python-docx, Pillow (beide bereits installiert)

Aufruf:
    from export_docx import generate_report
    buf = generate_report(result_dict, image_bytes, progress_cb=None)
    # buf ist ein BytesIO-Objekt → direkt als .docx herunterladen
"""

import io
import math
import base64
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

# ──────────────────────────────────────────────────────────────────────────────
# Konfiguration
# ──────────────────────────────────────────────────────────────────────────────

CROP_SIZE_PX   = 220     # Quadratischer Ausschnitt um das Objekt (Pixel im Originalbild)
CROP_SIZE_DOC  = 3.2     # Größe im Dokument (cm)
MAX_OBJECTS    = 500     # Sicherheitslimit (mehr = sehr langes Dokument)
TYPE_ICONS = {
    "galaxy":   "⬟",
    "star":     "★",
    "quasar":   "◎",
    "nebula":   "☁",
    "cluster":  "✦",
    "variable": "△",
    "double":   "⊕",
    "unknown":  "?",
}
TYPE_COLORS = {
    "galaxy":  RGBColor(0x5a, 0xab, 0xff),
    "star":    RGBColor(0xfd, 0xe0, 0x47),
    "quasar":  RGBColor(0xf4, 0x72, 0xb6),
    "nebula":  RGBColor(0x4a, 0xde, 0x80),
    "cluster": RGBColor(0xfb, 0x92, 0x3c),
}
COLOR_DARK  = RGBColor(0x1e, 0x1e, 0x2e)
COLOR_GRAY  = RGBColor(0x60, 0x60, 0x80)
COLOR_WHITE = RGBColor(0xff, 0xff, 0xff)
COLOR_ACCENT= RGBColor(0x2e, 0x75, 0xb6)


# ──────────────────────────────────────────────────────────────────────────────
# Hilfsfunktionen
# ──────────────────────────────────────────────────────────────────────────────

def _ra_str(ra_deg):
    h = int(ra_deg / 15)
    m = int((ra_deg / 15 - h) * 60)
    s = ((ra_deg / 15 - h) * 60 - m) * 60
    return f"{h:02d}h {m:02d}m {s:05.2f}s"


def _dec_str(dec_deg):
    sign = "+" if dec_deg >= 0 else "−"
    ad = abs(dec_deg)
    d = int(ad)
    m = int((ad - d) * 60)
    s = ((ad - d) * 60 - m) * 60
    return f"{sign}{d:02d}° {m:02d}' {s:04.1f}\""


def _fmt_dist(mly):
    if mly is None:
        return None
    if mly >= 1000:
        return f"{mly/1000:.2f} Mrd Lj"
    return f"{mly:.0f} Mio Lj"


def _cosmo(z):
    """
    Λ CDM Entfernungsberechnung (Planck 2018: H₀=67.4, Ωm=0.315, ΩΛ=0.685).
    Numerische Integration via Simpson-Regel (N=200 Schritte).
    Rückgabe: (dc_mly, da_mly, tl_gyr, age_at_emission_gyr)
    """
    if z is None or z <= 0:
        return None, None, None, None
    H0, Om, OL, c = 67.4, 0.315, 0.685, 299792.458
    Mpc_mly = 3.26156          # 1 Mpc = 3.26156 Mio Lj (exakt: 3.261563...)
    Mpc_m   = 3.08568e22       # 1 Mpc in Metern (korrekt: 3.085677581e22 m)
    Gyr_s   = 3.15576e16       # 1 Gyr in Sekunden (365.25 d/Jahr × 1e9)
    age_today_gyr = 13.787     # Planck 2018: 13.787 ± 0.020 Gyr

    def Ez(zz): return math.sqrt(Om * (1+zz)**3 + OL)
    N, dz = 200, z / 200
    dc = tl = 0
    for i in range(N+1):
        w  = 1 if i in (0, N) else (2 if i % 2 == 0 else 4)
        zi = i * dz
        e  = Ez(zi)
        dc += w / e
        tl += w / ((1+zi) * e)
    dc *= dz / 3
    tl *= dz / 3

    dH      = c / H0                        # Hubble-Distanz in Mpc
    H0_gyr  = (H0 * 1e3 / Mpc_m) * Gyr_s  # H₀ in 1/Gyr (korrekte Einheit)
    dc_mly  = dH * dc * Mpc_mly            # Komovingdistanz in Mio Lj
    da_mly  = dc_mly / (1 + z)             # Angulardistanz
    tl_gyr  = tl / H0_gyr                  # Lichtlaufzeit in Gyr
    age_emit= age_today_gyr - tl_gyr       # Alter des Universums beim Aussenden

    return dc_mly, da_mly, tl_gyr, age_emit


def _crop_object(img_pil, x, y, size=CROP_SIZE_PX):
    """Schneidet einen quadratischen Ausschnitt um (x, y) aus."""
    w, h = img_pil.size
    half = size // 2
    x0 = max(0, int(x) - half)
    y0 = max(0, int(y) - half)
    x1 = min(w, x0 + size)
    y1 = min(h, y0 + size)
    # Korrigiere wenn Rand erreicht
    x0 = max(0, x1 - size)
    y0 = max(0, y1 - size)
    crop = img_pil.crop((x0, y0, x1, y1))
    # Zielgröße normieren (falls Rand)
    if crop.size != (size, size):
        bg = Image.new("RGB", (size, size), (10, 10, 20))
        bg.paste(crop, (0, 0))
        crop = bg
    return crop


def _img_to_bytes(pil_img, fmt="PNG"):
    buf = io.BytesIO()
    pil_img.save(buf, format=fmt)
    buf.seek(0)
    return buf


def _set_cell_bg(cell, hex_color):
    """Setzt Hintergrundfarbe einer Tabellenzelle."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _add_run(para, text, bold=False, italic=False, size_pt=None, color=None, font=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if size_pt:
        run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    if font:
        run.font.name = font
    return run


def _hr(doc):
    """Horizontale Linie als Paragraph-Border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '3355aa')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ──────────────────────────────────────────────────────────────────────────────
# Haupt-Funktion
# ──────────────────────────────────────────────────────────────────────────────

def generate_report(result: dict, image_bytes: bytes,
                    progress_cb=None, max_objects=MAX_OBJECTS) -> io.BytesIO:
    """
    Erzeugt ein Word-Dokument mit Bild-Ausschnitt und Daten für jedes Objekt.

    result      : Das solve_field()-Ergebnis-Dict (enthält 'objects', 'wcs_info', ...)
    image_bytes : Das Originalbild als Bytes (JPEG oder PNG)
    progress_cb : Optionale Fortschritts-Callback-Funktion
    Rückgabe    : BytesIO mit dem .docx-Inhalt
    """
    def prog(m):
        if progress_cb: progress_cb(m)

    objects = result.get("objects", [])
    wcs     = result.get("wcs_info", {})
    n_total = len(objects)

    prog(f"[DOCX] Starte Export: {n_total} Objekte")

    # Bild laden
    try:
        img_pil = Image.open(io.BytesIO(image_bytes)).convert("RGB")
        img_w, img_h = img_pil.size
        has_image = True
    except Exception as e:
        prog(f"[DOCX] Bild nicht ladbar: {e}")
        img_pil = None
        img_w = img_h = 0
        has_image = False

    # Objekte sortieren: zuerst nach Typ-Priorität, dann nach Magnitude
    TYPE_ORDER = {"galaxy": 0, "quasar": 1, "cluster": 2, "nebula": 3,
                  "variable": 4, "double": 5, "star": 6, "unknown": 7}
    objs_sorted = sorted(
        [o for o in objects if o.get("x") is not None],
        key=lambda o: (TYPE_ORDER.get(o.get("type", "unknown"), 9),
                       float(o.get("magnitude") or 99))
    )
    if len(objs_sorted) > max_objects:
        prog(f"[DOCX] Auf {max_objects} Objekte begrenzt (von {len(objs_sorted)})")
        objs_sorted = objs_sorted[:max_objects]

    # Vollbild-Thumbnail für Deckblatt
    thumb_buf = None
    if has_image:
        thumb = img_pil.copy()
        thumb.thumbnail((600, 400), Image.LANCZOS)
        thumb_buf = _img_to_bytes(thumb)

    # ── Dokument aufbauen ─────────────────────────────────────────────────────
    doc = Document()

    # Seitenränder
    for section in doc.sections:
        section.page_width  = Cm(21)    # A4
        section.page_height = Cm(29.7)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    # Standard-Schrift
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10)

    # ── Deckblatt ─────────────────────────────────────────────────────────────
    # Titel
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(30)
    _add_run(title_p, "Astrometrie-Bericht", bold=True, size_pt=26,
             color=COLOR_ACCENT, font="Arial")

    # Untertitel
    ra_c  = wcs.get("ra_center")
    dec_c = wcs.get("dec_center")
    if ra_c is not None:
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(sub_p, f"Bildfeld: {_ra_str(ra_c)}  {_dec_str(dec_c)}",
                 size_pt=12, color=COLOR_GRAY)

    # Datum
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(date_p, datetime.now().strftime("%d. %B %Y"),
             size_pt=10, color=COLOR_GRAY, italic=True)

    # Statistik-Zeile
    stat_p = doc.add_paragraph()
    stat_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stat_p.paragraph_format.space_before = Pt(8)
    type_counts = {}
    for o in objects:
        t = o.get("type", "unknown")
        type_counts[t] = type_counts.get(t, 0) + 1
    stat_parts = [f"{n_total} Objekte erkannt"]
    for t in ["galaxy", "star", "quasar", "cluster", "nebula"]:
        if t in type_counts:
            stat_parts.append(f"{type_counts[t]} {t.capitalize()}n")
    _add_run(stat_p, "  ·  ".join(stat_parts), size_pt=10, color=COLOR_GRAY)

    # WCS-Info
    if wcs.get("scale_deg_per_px"):
        wcs_p = doc.add_paragraph()
        wcs_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        scale_arcsec = wcs["scale_deg_per_px"] * 3600
        _add_run(wcs_p,
                 f"Skala: {scale_arcsec:.2f}\"/px  ·  "
                 f"Rotation: {wcs.get('rotation_deg', 0):.1f}°  ·  "
                 f"Quelle: {wcs.get('source', 'ASTAP')}",
                 size_pt=9, color=COLOR_GRAY, italic=True)

    # Vollbild
    if thumb_buf:
        img_p = doc.add_paragraph()
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_p.paragraph_format.space_before = Pt(16)
        run = img_p.add_run()
        run.add_picture(thumb_buf, width=Cm(14))

    doc.add_page_break()

    # ── Objekte ───────────────────────────────────────────────────────────────
    prog(f"[DOCX] Erzeuge {len(objs_sorted)} Objekt-Einträge ...")

    for idx, obj in enumerate(objs_sorted):
        if idx % 20 == 0:
            prog(f"[DOCX] {idx}/{len(objs_sorted)} ...")

        otype   = obj.get("type", "unknown")
        oid     = obj.get("id", f"Objekt {idx+1}")
        oname   = obj.get("name") or oid
        cat     = obj.get("catalog", "")
        mag     = obj.get("magnitude")
        ra      = obj.get("ra")
        dec     = obj.get("dec")
        desc    = obj.get("description", "")
        morph   = obj.get("morph_type") or obj.get("hubble_type", "")
        px      = obj.get("x")
        py      = obj.get("y")
        # x_frac/y_frac sind auflösungsunabhängig → bevorzugen
        x_frac  = obj.get("x_frac")
        y_frac  = obj.get("y_frac")
        if x_frac is not None and y_frac is not None and img_w and img_h:
            px = x_frac * img_w
            py = y_frac * img_h

        # Rotverschiebung / Entfernung
        z_spec  = obj.get("redshift_z")
        z_photo = obj.get("photoz")
        z_desc  = None
        if z_spec is None and z_photo is None:
            import re
            m = re.search(r'z\s*[=≈]\s*([0-9.]+)', desc or "")
            if m:
                z_desc = float(m.group(1))
        z_use   = z_spec if z_spec is not None else (z_photo if z_photo is not None else z_desc)
        dc_mly, da_mly, tl_gyr, age_emit = _cosmo(z_use)
        if dc_mly is None:
            dc_mly = da_mly = tl_gyr = age_emit = None

        icon = TYPE_ICONS.get(otype, "?")
        color = TYPE_COLORS.get(otype, COLOR_GRAY)

        # ── Objekt-Tabelle: [Bild | Daten] ──────────────────────────────────
        # Crop-Größe an Bildauflösung anpassen (ca. 4% der Bildbreite)
        crop_px = max(180, min(600, int(img_w * 0.04))) if img_w else CROP_SIZE_PX
        tbl = doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        tbl.style = 'Table Grid'
        tbl.columns[0].width = Cm(3.8)
        tbl.columns[1].width = Cm(13.2)

        cell_img  = tbl.cell(0, 0)
        cell_data = tbl.cell(0, 1)
        cell_img.vertical_alignment  = WD_ALIGN_VERTICAL.TOP
        cell_data.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # Hintergrund der Bild-Zelle dunkel
        _set_cell_bg(cell_img, "0d0d1a")

        # ── Bild-Zelle ───────────────────────────────────────────────────────
        if has_image and px is not None and py is not None:
            try:
                crop = _crop_object(img_pil, px, py, crop_px)
                # Kreuz-Fadenkreuz einzeichnen
                from PIL import ImageDraw
                draw = ImageDraw.Draw(crop)
                cx_c, cy_c = crop_px // 2, crop_px // 2
                lc = (255, 180, 50, 200)
                gap = max(10, crop_px // 18)
                arm = max(25, crop_px // 6)
                draw.line([(cx_c - arm, cy_c), (cx_c - gap, cy_c)], fill=lc, width=1)
                draw.line([(cx_c + gap, cy_c), (cx_c + arm, cy_c)], fill=lc, width=1)
                draw.line([(cx_c, cy_c - arm), (cx_c, cy_c - gap)], fill=lc, width=1)
                draw.line([(cx_c, cy_c + gap), (cx_c, cy_c + arm)], fill=lc, width=1)
                crop_buf = _img_to_bytes(crop)
                p_img = cell_img.paragraphs[0]
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(4)
                run = p_img.add_run()
                run.add_picture(crop_buf, width=Cm(CROP_SIZE_DOC))
            except Exception as e:
                cell_img.paragraphs[0].add_run(f"[{e}]")
        else:
            p_no = cell_img.paragraphs[0]
            p_no.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_run(p_no, icon, size_pt=28, color=color)

        # ── Daten-Zelle ──────────────────────────────────────────────────────
        # Objekt-Name / Typ
        p_head = cell_data.paragraphs[0]
        p_head.paragraph_format.space_before = Pt(2)
        _add_run(p_head, f"{icon} ", size_pt=13, color=color, bold=True)
        _add_run(p_head, oname, bold=True, size_pt=13, color=RGBColor(0x22,0x22,0x44))
        if cat:
            _add_run(p_head, f"  [{cat}]", size_pt=9, color=COLOR_GRAY, italic=True)

        # Typ-Badge
        p_type = cell_data.add_paragraph()
        p_type.paragraph_format.space_before = Pt(1)
        _add_run(p_type, otype.upper(), bold=True, size_pt=8, color=color)
        if morph:
            _add_run(p_type, f"  ·  {morph}", size_pt=8, color=COLOR_GRAY)

        # Koordinaten
        if ra is not None:
            p_coord = cell_data.add_paragraph()
            p_coord.paragraph_format.space_before = Pt(4)
            _add_run(p_coord, "RA  ", bold=True, size_pt=9, color=COLOR_GRAY)
            _add_run(p_coord, _ra_str(ra), size_pt=9)
            _add_run(p_coord, "   Dec  ", bold=True, size_pt=9, color=COLOR_GRAY)
            _add_run(p_coord, _dec_str(dec), size_pt=9)

        # Magnitude
        if mag is not None:
            p_mag = cell_data.add_paragraph()
            p_mag.paragraph_format.space_before = Pt(2)
            _add_run(p_mag, "Magnitude  ", bold=True, size_pt=9, color=COLOR_GRAY)
            _add_run(p_mag, f"{float(mag):.2f} mag", size_pt=9)

        # SDSS Photometrie
        if obj.get("sdss_mag_r") is not None:
            bands = []
            for b, k in [("u","sdss_mag_u"),("g","sdss_mag_g"),("r","sdss_mag_r"),
                          ("i","sdss_mag_i"),("z","sdss_mag_z")]:
                v = obj.get(k)
                if v is not None:
                    bands.append(f"{b}={float(v):.2f}")
            if bands:
                p_sdss = cell_data.add_paragraph()
                p_sdss.paragraph_format.space_before = Pt(2)
                _add_run(p_sdss, "SDSS ugriz  ", bold=True, size_pt=9, color=COLOR_GRAY)
                _add_run(p_sdss, "  ".join(bands), size_pt=9)

        # Rotverschiebung + Entfernung
        if z_use is not None:
            p_z = cell_data.add_paragraph()
            p_z.paragraph_format.space_before = Pt(2)
            z_label = "z (spec)" if z_spec is not None else ("z (photo)" if z_photo is not None else "z")
            _add_run(p_z, f"{z_label}  ", bold=True, size_pt=9, color=COLOR_GRAY)
            _add_run(p_z, f"{z_use:.5f}", size_pt=9)
            cz = z_use * 299792
            _add_run(p_z, f"   →  {cz:,.0f} km/s", size_pt=9, color=COLOR_GRAY)

        if dc_mly is not None:
            p_dist = cell_data.add_paragraph()
            p_dist.paragraph_format.space_before = Pt(2)
            _add_run(p_dist, "Entfernung heute  ", bold=True, size_pt=9, color=COLOR_GRAY)
            _add_run(p_dist, _fmt_dist(dc_mly), size_pt=9)

            p_tl = cell_data.add_paragraph()
            p_tl.paragraph_format.space_before = Pt(1)
            _add_run(p_tl, "Licht gesendet vor  ", bold=True, size_pt=9, color=COLOR_GRAY)
            _add_run(p_tl, f"{tl_gyr:.2f} Mrd Jahren", size_pt=9)
            if age_emit is not None:
                _add_run(p_tl,
                         f"  (Universum war {age_emit:.2f} Mrd J. alt, "
                         f"Distanz damals: {_fmt_dist(da_mly)})",
                         size_pt=8, color=COLOR_GRAY)

        # Beschreibung
        if desc and desc != "SDSS PhotoObj":
            # Nur sinnvolle Beschreibungen zeigen (nicht reine Magnitude-Listen)
            clean_desc = desc.replace(" | ", "  ·  ")
            if len(clean_desc) > 5:
                p_desc = cell_data.add_paragraph()
                p_desc.paragraph_format.space_before = Pt(3)
                _add_run(p_desc, clean_desc, size_pt=8, color=COLOR_GRAY, italic=True)

        # Pixel-Position
        if px is not None:
            p_px = cell_data.add_paragraph()
            p_px.paragraph_format.space_before = Pt(2)
            _add_run(p_px, "Position  ", bold=True, size_pt=8, color=COLOR_GRAY)
            _add_run(p_px, f"x={int(px)}  y={int(py)} px", size_pt=8, color=COLOR_GRAY)

        # Trennlinie nach jedem Objekt
        doc.add_paragraph().paragraph_format.space_before = Pt(3)
        _hr(doc)
        doc.add_paragraph().paragraph_format.space_before = Pt(3)

    # ── Letzte Seite: Zusammenfassung ─────────────────────────────────────────
    doc.add_page_break()
    sum_p = doc.add_paragraph()
    _add_run(sum_p, "Zusammenfassung", bold=True, size_pt=16, color=COLOR_ACCENT)

    stat_tbl = doc.add_table(rows=len(type_counts)+1, cols=2)
    stat_tbl.style = 'Table Grid'
    stat_tbl.columns[0].width = Cm(8)
    stat_tbl.columns[1].width = Cm(4)

    # Header
    _set_cell_bg(stat_tbl.cell(0,0), "2e75b6")
    _set_cell_bg(stat_tbl.cell(0,1), "2e75b6")
    h0 = stat_tbl.cell(0,0).paragraphs[0]
    h1 = stat_tbl.cell(0,1).paragraphs[0]
    _add_run(h0, "Objekttyp", bold=True, size_pt=10, color=COLOR_WHITE)
    _add_run(h1, "Anzahl",    bold=True, size_pt=10, color=COLOR_WHITE)

    for i, (t, n) in enumerate(sorted(type_counts.items(), key=lambda x: -x[1])):
        row = stat_tbl.rows[i+1]
        c0, c1 = row.cells[0], row.cells[1]
        p0 = c0.paragraphs[0]
        p1 = c1.paragraphs[0]
        icon = TYPE_ICONS.get(t, "?")
        col  = TYPE_COLORS.get(t, COLOR_GRAY)
        _add_run(p0, f"{icon} {t.capitalize()}", size_pt=10, color=col)
        _add_run(p1, str(n), size_pt=10, bold=True)

    # Footer-Info
    foot_p = doc.add_paragraph()
    foot_p.paragraph_format.space_before = Pt(20)
    _add_run(foot_p,
             f"Erstellt mit Plate Solver  ·  {datetime.now().strftime('%d.%m.%Y %H:%M')}  ·  "
             f"Kosmologie: H₀=67.4 km/s/Mpc, Ωm=0.315 (Planck 2018)",
             size_pt=8, color=COLOR_GRAY, italic=True)

    prog(f"[DOCX] Dokument fertig, packe ...")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    prog(f"[DOCX] Export abgeschlossen ({buf.getbuffer().nbytes // 1024} KB)")
    return buf
