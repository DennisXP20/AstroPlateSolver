"""Erzeugt Dokumentation.docx mit Feature-Beschreibung und Wissenschafts-Hintergrund."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Styles
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

def h1(text):
    p = doc.add_heading(text, level=1)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    return p

def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p

def bullet(text):
    doc.add_paragraph(text, style="List Bullet")

def code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)

# ================================================================
# TITEL
# ================================================================
title = doc.add_heading("Astro Plate Solver", 0)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Dokumentation · Features und wissenschaftlicher Hintergrund · v1.0-rc1")
r.italic = True
r.font.size = Pt(11)
doc.add_paragraph()

# ================================================================
# EINLEITUNG
# ================================================================
h1("1. Überblick")
para(
    "Der Astro Plate Solver ist eine lokale Webanwendung zur astrometrischen "
    "Analyse von Astrofotos. Sie erkennt automatisch, welcher Himmelsausschnitt "
    "auf einem Bild zu sehen ist (Plate Solving), identifiziert sichtbare Objekte "
    "aus einem lokalen Katalog, führt differentielle Aperturphotometrie durch "
    "und visualisiert die Galaxienverteilung auf der Hubble-Sequenz."
)
para(
    "Die Anwendung besteht aus einem Python-Backend (HTTP-Server plus wissenschaftliche "
    "Module) und einer Single-Page-Web-Oberfläche. Nach einem einmaligen Katalog-Download "
    "läuft sie komplett offline — nur optionale SIMBAD-Lookups für Galaxien-Morphologie "
    "benötigen Internet."
)

h2("1.1 Zielgruppe")
bullet("Amateur-Astrofotografen, die wissen möchten, was auf ihren Bildern zu sehen ist.")
bullet("Schüler und Studierende, die Photometrie und WCS-Mathematik lernen.")
bullet("Hobbyforscher, die Helligkeiten oder Galaxienklassen aus eigenen Aufnahmen gewinnen wollen.")

# ================================================================
# ARCHITEKTUR
# ================================================================
h1("2. Architektur")
para(
    "Die Anwendung ist als einfacher HTTP-Server (stdlib) plus statische HTML-Oberfläche "
    "aufgebaut. Jede wissenschaftliche Funktion ist in einem eigenen Python-Modul gekapselt, "
    "das bei Bedarf lazy geladen wird. Das Frontend kommuniziert ausschließlich über JSON-APIs."
)

h2("2.1 Module")
bullet("server.py — HTTP-Server, API-Endpunkte, Hintergrundthreads für Downloads und Batch-Jobs.")
bullet("solver.py — Kern: Bild-IO, Sternextraktion, Pattern-Matching, WCS-Aufbau, Katalogabfrage, ASTAP-Integration.")
bullet("catalog_dl.py — Download von Tycho-2, Gaia DR3, Messier/NGC/IC von VizieR und Wikidata.")
bullet("import_catalogs.py — Kommandozeilen-Werkzeug für manuelle CSV-Imports (PGC, Quasare).")
bullet("photometry.py — Aperturphotometrie mit Ringhintergrund, sRGB-Linearisierung, Zeropoint-Fit pro Katalog.")
bullet("galaxy_morph.py — SIMBAD-TAP-Abfrage für Galaxien-Morphologietypen mit lokalem Cache.")
bullet("index.html — Single-Page-App mit Tabs (Solve, Fotometrie, Objekte, Statistik, Hilfe).")

h2("2.2 Datenfluss beim Solven")
bullet("Browser lädt Bild per Base64 hoch (/api/solve).")
bullet("Server entpackt das Bild, startet ASTAP (Subprozess).")
bullet("ASTAP schreibt eine .wcs- oder .ini-Datei mit Referenzpixel und CD-Matrix.")
bullet("solver.py extrahiert Sterne, baut ein TanWCS-Objekt und fragt den lokalen Katalog ab.")
bullet("Jedes Katalogobjekt bekommt Pixel-Koordinaten; Zurückgabe als JSON-Liste.")
bullet("Frontend zeichnet die Overlays ins Canvas, berechnet Statistiken und Hubble-Sequenz.")

# ================================================================
# FEATURES
# ================================================================
h1("3. Features im Detail")

h2("3.1 Plate Solving")
para(
    "Das Plate Solving ist die Kernfunktion: aus einem Bild ohne Metadaten wird die "
    "Himmelsposition, der Bildmaßstab, die Rotation und die volle CD-Matrix abgeleitet."
)
h3("Methode")
bullet(
    "Primärer Solver ist ASTAP (externes Programm). ASTAP nutzt geometrische "
    "Dreieckmusterabgleiche zwischen Bildsternen und einem Sternkatalog (D80, H18 o. ä.)."
)
bullet(
    "Der Solver ruft ASTAP mit vollem Himmelsradius (-r 180) auf und probiert "
    "automatisch mehrere Gesichtsfelder (-fov 2/5/15 als Fallback), wenn der erste "
    "Versuch scheitert."
)
bullet(
    "Nach erfolgreichem Lauf liest solver.py die von ASTAP geschriebene .wcs- oder "
    ".ini-Datei. Die CD-Matrix (CD1_1, CD1_2, CD2_1, CD2_2) beschreibt die vollständige "
    "lineare Abbildung zwischen Pixel- und Weltkoordinaten."
)
h3("Koordinatentransformation")
para(
    "Die Umrechnung zwischen Pixeln und Himmelskoordinaten erfolgt über die TAN-Projektion "
    "(Gnomonisch). Das ist die Standardprojektion der WCS-FITS-Norm und für typische "
    "Teleskop-Bildwinkel eine sehr gute Näherung. Der Algorithmus:"
)
bullet("Pixel (x,y) → Versatz zum Referenzpixel (dx,dy).")
bullet("Standard-Koordinaten (xi, eta) = CD · (dx, dy).")
bullet("Deprojektion: (xi, eta) + Tangentialpunkt → (RA, Dec).")
bullet("Rückweg invers: TanWCS.world_to_pixel() für jedes Katalogobjekt.")

h2("3.2 Sternextraktion")
para(
    "Für Bilder ohne Katalog-Fit (manueller WCS-Modus) und für die Photometrie-Detektion "
    "wird eine eigene Sternextraktion genutzt:"
)
bullet("Hintergrund-Schätzung per 25. Perzentil, Log-Kompression für Dynamik.")
bullet("Schwellwert-Detektion, verbundene Komponenten, Peak-Finding.")
bullet("Flussgewichtete Schwerpunktbestimmung (Sub-Pixel-Genauigkeit).")
bullet("Maximale Sternzahl konfigurierbar (Standard 300–800).")

h2("3.3 Objektkatalog und Matching")
para("Der lokale Katalog enthält:")
bullet("Messier (110 Einträge), NGC und IC (ca. 13 000).")
bullet("PGC-Galaxien (ca. 1 Mio., via beigelegter CSV).")
bullet("Tycho-2 Sterne bis mag 12 (ca. 2,5 Mio.).")
bullet("Gaia DR3 Sterne bis mag 17 (nach Download-Auswahl).")
bullet("Quasar-Katalog (ca. 500 000, aus Milliquas-Subset).")
para(
    "Beim Solven fragt der Server alle Kataloge ab, die im Sucherechteck um den "
    "Bildmittelpunkt liegen und den mag-Limit-Filter erfüllen. Jeder Treffer wird "
    "mit world_to_pixel auf Bildkoordinaten gebracht; nur Objekte innerhalb der "
    "Bildgrenzen werden zurückgegeben."
)

h2("3.4 Aperturphotometrie")
para(
    "Aperturphotometrie misst die Helligkeit eines Sterns, indem der gesamte Fluss "
    "innerhalb eines Kreises (Apertur) aufsummiert und der Himmelshintergrund aus einem "
    "umgebenden Ring abgezogen wird. Die Implementierung in photometry.py folgt dem "
    "klassischen Schema:"
)
h3("Ablauf")
bullet("Referenzsterne aus dem Solve-Ergebnis übernehmen (Tycho-2, Gaia DR3).")
bullet("Jeden Stern um lokales Pixelmaximum verfeinern (±8 px Suchfenster).")
bullet("Erster Messdurchgang mit vom Nutzer gesetzter Apertur (Standard r_ap=5).")
bullet("FWHM aus der 2D-Varianz der Apertur-Verteilung schätzen.")
bullet("Adaptive Apertur: r_ap = max(user, 1.5·FWHM_med), r_in = r_ap+3, r_out = r_ap+7.")
bullet("Zweiter Messdurchgang mit adaptiver Apertur (konsistente PSF-Erfassung).")
bullet("Sättigungs-Filter: Sterne mit Peak ≥ 98 % des Bildmaximums werden verworfen.")

h3("Formeln")
para("Fluss mit Ringhintergrund:")
code("F  =  Σ (pixel in aperture)  −  bg · N_aperture")
code("bg =  sigma-geclipter Median der Ringpixel")
para("Instrumentelle Magnitude:")
code("m_inst = −2,5 · log₁₀(F)")
para("Zeropoint-Bestimmung (robust, sigma-clipped Median):")
code("ZP        = Median(m_kat − m_inst)")
code("σ(ZP)     = 1,4826 · MAD(m_kat − m_inst)")
code("m_kalib   = m_inst + ZP")
para("Signal-Rausch-Verhältnis nach CCD-Gleichung:")
code("σ²  =  F_src  +  N_ap · σ_bg²  +  (N_ap² · σ_bg²) / N_ring")
code("SNR =  F_src / σ")
para(
    "Der Faktor 1,4826 wandelt den Median-Absolutwert (MAD) in ein Gauß-σ-Äquivalent um — "
    "robust gegen Ausreißer, anders als eine reine Standardabweichung."
)

h3("sRGB-Linearisierung")
para(
    "JPEG- und PNG-Bilder werden mit sRGB-Gamma (γ ≈ 2,2) kodiert, was Helligkeiten "
    "nichtlinear verzerrt. Ohne Korrektur wäre der gemessene Fluss systematisch falsch. "
    "photometry.load_raw() wendet die inverse sRGB-Transformation (IEC 61966-2-1) "
    "kanalweise auf R, G, B an und bildet eine Rec.709-gewichtete Luminanzkarte:"
)
code("L = 0,2126·R_lin + 0,7152·G_lin + 0,0722·B_lin")
para(
    "FITS-Dateien bleiben unverändert, da sie bereits lineare ADU enthalten."
)

h3("Band-Trennung")
para(
    "Tycho-2 liefert V-Band-Magnituden (Vt), Gaia DR3 liefert G-Band (breites optisches "
    "Band). Ohne Farbterm-Korrektur würde das Mischen beider Bänder systematische "
    "Offsets von 0,1–0,3 mag erzeugen. Die Kalibrierung berechnet deshalb für jeden "
    "Katalog einen eigenen Zeropoint und benutzt für die Ausgabe den primären Band "
    "(bevorzugt Gaia DR3). Jeder Referenzstern wird gegen den Zeropoint seines "
    "eigenen Katalogs verrechnet."
)

h3("Qualitätsmetriken")
para("Jedes Photometrie-Ergebnis enthält diagnostische Felder:")
bullet("zp_per_catalog — Zeropoint und Streuung pro Band.")
bullet("primary_band — welcher Katalog als Referenz verwendet wurde.")
bullet("zp_std / robust_sigma — Zeropoint-Streuung; < 0,1 mag = exzellent.")
bullet("median_residual — systematischer Offset; sollte nahe 0 sein.")
bullet("rms — quadratischer Mittelwert der Residuen im Primärband.")
bullet("quality — qualitative Einordnung: exzellent / gut / brauchbar / schwach.")
bullet("n_saturated_skipped — ausgeschlossene Sättigungs-Kandidaten.")
bullet("mean_fwhm — Median der FWHM über alle Referenzsterne.")

h2("3.5 Hubble-Sequenz")
para(
    "Die Hubble-Sequenz (Stimmgabeldiagramm, Edwin Hubble 1926) ist das klassische "
    "Schema zur Klassifikation von Galaxien nach ihrer visuellen Morphologie: "
    "Elliptische Galaxien (E0–E7) links, Lentikuläre (S0) am Knoten, "
    "normale Spiralen (SAa–SAd) auf dem oberen Arm, Balkenspiralen (SBa–SBd) "
    "auf dem unteren, und Irreguläre ganz rechts. Die Einordnung ist rein visuell "
    "und beschreibt keine zeitliche Entwicklung."
)
h3("Datenquellen")
bullet("Hardcoded NED/RC3-Referenz für bekannte Messier- und NGC-Galaxien (ca. 70 Einträge).")
bullet("Regex-Parser für Morphologie-Strings aus der eigenen Objektbeschreibung.")
bullet("Online-Fallback: SIMBAD TAP/ADQL-Abfrage des morph_type-Feldes.")
bullet("Lokaler Cache in morph_cache.db und browserseitig in localStorage.")
h3("Parser")
para(
    "Der Morphologie-Parser in classifyFromMorphString() toleriert die üblichen "
    "Notationsvarianten: Ring-Präfixe (R), (R'), (r), (s), (rs); Caret-Zusätze wie "
    "SB0^0^; Peculiarity-Marker (pec, :); Unsicherheitsmarker (?); nackte Spiralen "
    "ohne SA/SB-Präfix (Sa, Sb, Sc, Sd, Sdm, Sm) und lentikuläre Varianten. "
    "Jeder Typ wird auf eine numerische Position 0…10 abgebildet (E0 links, Irr rechts)."
)

h2("3.6 Weitere Funktionen")
bullet("Batch-Modus: Mehrere Bilder in einem Rutsch solven, Ergebnisse als JSON.")
bullet("FOV-Vergleich: Visualisierung des Bildfeldes gegen bekannte Objekte (Mond, M31 …).")
bullet("Statistik-Panel: Magnitude-Histogramm, Katalogverteilung, Hubble-Diagramm.")
bullet("Manueller WCS-Modus: bekannte Koordinaten eingeben statt solving.")
bullet("FITS-Import mit Header-Parser (CRVAL, CRPIX, CDELT, CROTA).")

# ================================================================
# WISSENSCHAFTLICHER HINTERGRUND
# ================================================================
h1("4. Wissenschaftlicher Hintergrund")

h2("4.1 Plate Solving")
para(
    "Plate Solving bedeutet: das Programm erkennt ein Bildmuster von Sternen und findet "
    "dazu die passende Himmelsposition, indem es Dreiecke oder Quads aus den hellsten "
    "Sternen bildet und diese Formen in einem vorab indizierten Sternkatalog sucht. "
    "Das Verfahren ist invariant gegen Rotation und Skalierung und funktioniert auch "
    "ganz ohne Vorwissen über die Aufnahmeparameter."
)
para(
    "ASTAP verwendet eine eigene Implementierung dieses Prinzips mit tiefen Katalogen "
    "(D80, H18, V17, G17), die von der ASTAP-Webseite bezogen werden. Der Tiefengrad "
    "(mag 8 bis mag 18) bestimmt, welche Bildfelder gelöst werden können: für weite "
    "Felder (> 3°) reicht D80, für Teleskop-Nahaufnahmen braucht man H17/H18."
)

h2("4.2 Magnitudensystem")
para(
    "Die scheinbare Magnitude ist ein logarithmisches Maß für die Helligkeit eines "
    "Objekts von der Erde aus gesehen. Sie ist umgekehrt skaliert: kleinere Zahl = heller. "
    "Der Zusammenhang mit dem Fluss ist:"
)
code("m₁ − m₂ = −2,5 · log₁₀(F₁ / F₂)")
para(
    "Eine Differenz von 5 mag entspricht einem Helligkeitsverhältnis von 100. "
    "Referenzwerte: Sonne −26,7; Vollmond −12,6; Venus bis −4,9; Sirius −1,46; "
    "Auge-Grenze +6; Hubble-Grenze ≈ +31."
)
para(
    "Der Zeropoint eines Bildes ist die Konstante, die die instrumentelle Skala "
    "(−2,5·log F in ADU) mit der physikalischen kat-Magnitude verbindet. Er hängt von "
    "Teleskopapertur, Belichtungszeit, Filter, Himmelstransparenz und Luftmasse ab — "
    "und wird deshalb für jedes Bild neu bestimmt."
)

h2("4.3 PSF und FWHM")
para(
    "Die Point Spread Function (PSF) beschreibt, wie ein punktförmiger Stern vom "
    "optischen System auf dem Sensor verschmiert wird. Die Full Width at Half Maximum "
    "(FWHM) ist der Durchmesser der PSF bei halbem Peak — ein Maß für die Schärfe "
    "des Bildes. Typische Werte: 1''–2'' bei professionellen Teleskopen unter gutem "
    "Seeing, 3''–5'' bei Amateur-Aufnahmen. Für gute Photometrie sollte die Apertur "
    "mindestens 1,5·FWHM groß sein, damit > 95 % des Sternflusses erfasst werden. "
    "Genau das macht die adaptive Apertur automatisch."
)

h2("4.4 Hubble-Sequenz")
para(
    "Hubbles Diagramm wurde 1926 als rein beschreibendes Schema eingeführt, wurde "
    "aber lange als Evolutionsfolge missverstanden (die Begriffe 'frueh' und 'spaet' "
    "Galaxien stammen daher). Die moderne Astrophysik sieht die Morphologie als "
    "Folge von Umgebungsdichte, Verschmelzungsgeschichte und Kühlungseffizienz, "
    "nicht als Zeitachse. Das Diagramm bleibt aber ein unverzichtbares Klassifikations-"
    "und Kommunikationswerkzeug."
)
para("Positionen entlang der Sequenz (numerisch 0–10 intern):")
bullet("E0–E7 (0,0–2,2): elliptische Galaxien, Zahl = Elliptizität × 10.")
bullet("S0/SAB0/SB0 (2,8): lentikulär, Scheibe ohne Spiralarme.")
bullet("Sa–Sd (4,8–7,6): Spiralen mit abnehmendem Bulge und offener werdenden Armen.")
bullet("Sm (9,0): Magellansche Spirale, sehr späte, unregelmäßige Form.")
bullet("Irr (9,5): irreguläre Galaxien ohne erkennbare Struktur.")

# ================================================================
# BEDIENUNG
# ================================================================
h1("5. Bedienung")

h2("5.1 Plate Solve")
bullet("Bild ins Upload-Feld ziehen.")
bullet("mag-Limit wählen (Standard 19 reicht meist; höher = mehr Treffer, länger).")
bullet("Aktive Kataloge wählen (Tycho-2, Gaia DR3 …).")
bullet("\"Plate Solve\" klicken. Ein ASTAP-Fenster öffnet sich mit dem Fortschritt.")
bullet(
    "Nach Abschluss erscheinen RA, Dec, Bildmaßstab, Rotation und alle gefundenen "
    "Objekte als Overlay. Hover über einen Kreis zeigt Katalog-ID und Magnitude."
)

h2("5.2 Fotometrie")
bullet("Nach erfolgreichem Solve in den Tab Fotometrie wechseln.")
bullet("Apertur-Radien einstellen (Standard r_ap=5, r_in=8, r_out=12 Pixel).")
bullet("Mag-Limit für Referenzsterne festlegen (Standard 16).")
bullet("Referenzkataloge aktivieren (Gaia DR3 bevorzugt).")
bullet("Aperturphotometrie starten klicken.")
bullet(
    "Das Ergebnis enthält Zeropoint, RMS, Quality-Label und eine Tabelle mit "
    "jedem Referenzstern (klickbar → im Bild markiert)."
)

h2("5.3 Statistik und Hubble-Sequenz")
bullet("Panel Statistik oeffnen.")
bullet("Magnitude-Histogramm zeigt die Helligkeitsverteilung aller Objekte.")
bullet("Hubble-Stimmgabel zeigt klassifizierte Galaxien entlang der Sequenz.")
bullet(
    "Nicht klassifizierte Galaxien werden als farbkodierte Chips darunter aufgelistet. "
    "Ein Klick auf Online suchen (SIMBAD) fragt für jeden unklassifizierten "
    "Eintrag SIMBAD nach der Morphologie. Treffer werden lokal gecacht und "
    "klassifizieren die Galaxie beim nächsten Rendern automatisch."
)

# ================================================================
# GRENZEN
# ================================================================
h1("6. Grenzen und Hinweise")
bullet(
    "Photometrie auf JPEG/PNG ist nur eine Näherung: sRGB-Linearisierung korrigiert "
    "die Gamma-Kurve, aber JPEG-Kompressionsartefakte, Weißabgleich und lokale "
    "Kontrast-Anpassungen in der Kamera bleiben Fehlerquellen. Für wissenschaftlich "
    "saubere Messungen FITS-Rohdaten verwenden."
)
bullet(
    "Crowding in dichten Feldern (Kugelsternhaufen, Galaxienzentren) biasiert die "
    "Aperturphotometrie, weil Nachbarsterne in die Apertur leaken. Dann hilft nur "
    "PSF-Photometrie (nicht implementiert in v1)."
)
bullet(
    "Die Hubble-Klassifikation funktioniert nur für Galaxien, die SIMBAD im Morph-Feld "
    "kennt. Viele PGC-Galaxien sind dort ohne Morph-Eintrag — sie werden orange "
    "markiert (kein Morph-Typ hinterlegt)."
)
bullet(
    "Das Tool ist für Bildung und Amateurbeobachtung gedacht, nicht als Ersatz "
    "für professionelle Pipelines wie SExtractor, SCAMP, PSFEx oder DAOPHOT."
)

# Fußzeile
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Astro Plate Solver v1.0-rc1 · Dokumentation")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.save("Dokumentation.docx")
print("Dokumentation.docx erstellt")
